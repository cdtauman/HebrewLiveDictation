import os
import tempfile
import unittest

from hebrew_live_dictation import export, history
from hebrew_live_dictation.config import Config


class HistoryTests(unittest.TestCase):
    def test_append_and_load_roundtrip(self):
        with tempfile.TemporaryDirectory() as tmp:
            config = Config(tmp)
            self.assertTrue(history.append(config, "שלום עולם", target="notepad.exe", when=1000))
            entries = history.load(config)
            self.assertEqual(len(entries), 1)
            self.assertEqual(entries[0]["text"], "שלום עולם")
            self.assertEqual(entries[0]["target"], "notepad.exe")
            self.assertEqual(entries[0]["ts"], 1000)
            self.assertTrue(entries[0]["id"])

    def test_disabled_does_not_record(self):
        with tempfile.TemporaryDirectory() as tmp:
            config = Config(tmp)
            config.update({"history.enabled": False})
            self.assertFalse(history.append(config, "hi"))
            self.assertEqual(history.load(config), [])

    def test_delete_by_entry_id(self):
        with tempfile.TemporaryDirectory() as tmp:
            config = Config(tmp)
            history.append(config, "keep", when=1)
            history.append(config, "delete me", when=2)
            entries = history.load(config)
            self.assertTrue(history.delete(config, entries[1]["id"]))
            self.assertEqual([e["text"] for e in history.load(config)], ["keep"])

    def test_legacy_entry_id_is_stable(self):
        entry = {"ts": 7, "target": "notepad.exe", "text": "legacy"}
        self.assertEqual(history.entry_id(entry), history.entry_id(dict(entry)))
        normalized = history.normalize_entry(entry)
        self.assertEqual(normalized["text"], "legacy")
        self.assertEqual(normalized["chars"], 6)
        self.assertTrue(normalized["id"])

    def test_blank_text_ignored(self):
        with tempfile.TemporaryDirectory() as tmp:
            config = Config(tmp)
            self.assertFalse(history.append(config, "   "))

    def test_trim_to_max_entries(self):
        with tempfile.TemporaryDirectory() as tmp:
            config = Config(tmp)
            config.update({"history.max_entries": 3})
            for i in range(6):
                history.append(config, f"line {i}", when=i)
            entries = history.load(config)
            self.assertEqual([e["text"] for e in entries], ["line 3", "line 4", "line 5"])

    def test_clear(self):
        with tempfile.TemporaryDirectory() as tmp:
            config = Config(tmp)
            history.append(config, "x")
            self.assertTrue(history.clear(config))
            self.assertEqual(history.load(config), [])


class ExportTests(unittest.TestCase):
    def test_entries_to_text(self):
        text = export.entries_to_text([{"ts": 0, "target": "t", "text": "hello"}])
        self.assertIn("hello", text)
        self.assertIn("[t]", text)

    def test_write_txt(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "out.txt")
            export.write_txt(path, "שלום עולם")
            with open(path, "r", encoding="utf-8") as f:
                self.assertEqual(f.read(), "שלום עולם")

    def test_write_docx_is_rtl(self):
        try:
            import docx  # noqa: F401
        except Exception:
            self.skipTest("python-docx not installed")
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            path = os.path.join(tmp, "out.docx")
            export.write_docx(path, "שלום עולם\nשורה שנייה")
            doc = Document(path)
            texts = [p.text for p in doc.paragraphs]
            self.assertIn("שלום עולם", texts)
            # RTL marker present on the first paragraph's properties.
            xml = doc.paragraphs[0]._p.xml
            self.assertIn("bidi", xml)


if __name__ == "__main__":
    unittest.main()
